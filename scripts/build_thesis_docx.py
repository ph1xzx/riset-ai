from __future__ import annotations

import re
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

BASE = Path(__file__).resolve().parents[1]
OUT = BASE / "DRAFT-SKRIPSI-PT-PRISMA-GAPURA.docx"
CHAPTERS = [
    BASE / "DRAFT-BAB-I-PENDAHULUAN.md",
    BASE / "DRAFT-BAB-II-LANDASAN-TEORI.md",
    BASE / "DRAFT-BAB-III-IV-SPK-PRISMA-GAPURA.md",
    BASE / "DRAFT-BAB-V-PENUTUP.md",
    BASE / "DRAFT-DAFTAR-PUSTAKA.md",
    BASE / "DRAFT-LAMPIRAN.md",
]
FIGURES_DIR = BASE / "research_artifacts" / "screenshots" / "thesis"
UML_DIR = BASE / "research_artifacts" / "uml_export"
BLACK = RGBColor(0, 0, 0)

FIGURE_IMAGES = {
    "Gambar 3.1 Arsitektur Sistem SPK": UML_DIR / "08-arsitektur-sistem.png",
    "Gambar 3.2 Entity Relationship Diagram (ERD)": UML_DIR / "05-erd-database.png",
    "Gambar 3.3 Logical Record Structure (LRS)": UML_DIR / "09-lrs-database.png",
    "Gambar 3.4 Use Case Diagram": UML_DIR / "03-use-case-admin.png",
    "Gambar 3.5 Activity Diagram Perhitungan PROMETHEE II": UML_DIR / "04-activity-promethee.png",
    "Gambar 3.6 Sequence Diagram Proses Perhitungan": UML_DIR / "06-sequence-perhitungan.png",
    "Gambar 3.7 Class Diagram": UML_DIR / "07-class-diagram.png",
    "Gambar 3.8 Rancangan Landing Page Profile": BASE / "research_artifacts" / "mockups" / "01-mockup-landing-profile.png",
    "Gambar 3.9 Rancangan Login Admin": BASE / "research_artifacts" / "mockups" / "02-mockup-login-admin.png",
    "Gambar 3.10 Rancangan Dashboard Admin": BASE / "research_artifacts" / "mockups" / "03-mockup-dashboard.png",
    "Gambar 3.11 Rancangan Data Transaksi": BASE / "research_artifacts" / "mockups" / "04-mockup-data-transaksi.png",
    "Gambar 3.12 Rancangan Kriteria dan Bobot": BASE / "research_artifacts" / "mockups" / "05-mockup-kriteria-bobot.png",
    "Gambar 3.13 Rancangan Kelola Data": BASE / "research_artifacts" / "mockups" / "06-mockup-kelola-data.png",
    "Gambar 3.14 Rancangan Perhitungan PROMETHEE II": BASE / "research_artifacts" / "mockups" / "07-mockup-perhitungan.png",
    "Gambar 3.15 Rancangan Hasil Ranking": BASE / "research_artifacts" / "mockups" / "08-mockup-hasil-ranking.png",
    "Gambar 3.16 Rancangan Panduan": BASE / "research_artifacts" / "mockups" / "09-mockup-panduan.png", 
    "Gambar 4.1 Implementasi Landing Page Profile": FIGURES_DIR / "09-profile-landing.png",
    "Gambar 4.2 Implementasi Halaman Login Admin": FIGURES_DIR / "01-login-admin.png",
    "Gambar 4.3 Implementasi Dashboard Admin": FIGURES_DIR / "02-dashboard-admin.png",
    "Gambar 4.4 Implementasi Halaman Data Transaksi": FIGURES_DIR / "03-data-transaksi.png",
    "Gambar 4.5 Implementasi Halaman Kelola Data": FIGURES_DIR / "05-kelola-data.png",
    "Gambar 4.6 Implementasi Halaman Kriteria dan Bobot": FIGURES_DIR / "04-kriteria-bobot.png",
    "Gambar 4.7 Implementasi Halaman Perhitungan PROMETHEE II": FIGURES_DIR / "06-perhitungan-promethee.png",
    "Gambar 4.8 Implementasi Halaman Hasil Ranking": FIGURES_DIR / "07-hasil-ranking.png",
}


def set_cell_text(cell, text: str, bold: bool = False):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if bold else WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1
    clean_text = text.strip().replace('`', '').replace('**', '')
    r = p.add_run(clean_text)
    r.font.name = "Times New Roman"
    r.font.size = Pt(10)
    r.font.color.rgb = BLACK
    r.bold = bold


def add_page_number(paragraph, alignment=WD_ALIGN_PARAGRAPH.CENTER):
    paragraph.text = ""
    paragraph.alignment = alignment
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)
    run.font.name = "Times New Roman"
    run.font.size = Pt(10)
    run.font.color.rgb = BLACK


def _clear_header_footer(container):
    for paragraph in container.paragraphs:
        paragraph.text = ""
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT


def _set_page_number_type(section, fmt, start=None):
    sect_pr = section._sectPr
    for existing in list(sect_pr):
        if existing.tag == qn("w:pgNumType"):
            sect_pr.remove(existing)
    pg_num_type = OxmlElement("w:pgNumType")
    pg_num_type.set(qn("w:fmt"), fmt)
    if start is not None:
        pg_num_type.set(qn("w:start"), str(start))
    sect_pr.append(pg_num_type)


def _set_section_geometry(section):
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(3)
    section.bottom_margin = Cm(3)
    section.left_margin = Cm(4)
    section.right_margin = Cm(3)
    section.header_distance = Cm(1.25)
    section.footer_distance = Cm(1.25)


def configure_cover_section(section):
    """The cover is a standalone unnumbered section."""
    _set_section_geometry(section)
    section.different_first_page_header_footer = False
    section.header.is_linked_to_previous = False
    section.footer.is_linked_to_previous = False
    _clear_header_footer(section.header)
    _clear_header_footer(section.footer)


def configure_front_matter_section(section):
    """The preliminary matter uses lower Roman numerals, beginning at i."""
    _set_section_geometry(section)
    section.different_first_page_header_footer = True
    section.header.is_linked_to_previous = False
    section.footer.is_linked_to_previous = False
    section.first_page_header.is_linked_to_previous = False
    section.first_page_footer.is_linked_to_previous = False
    _clear_header_footer(section.header)
    _clear_header_footer(section.first_page_header)
    _clear_header_footer(section.first_page_footer)
    _clear_header_footer(section.footer)
    add_page_number(section.first_page_footer.paragraphs[0], WD_ALIGN_PARAGRAPH.CENTER)
    add_page_number(section.footer.paragraphs[0], WD_ALIGN_PARAGRAPH.CENTER)
    _set_page_number_type(section, "lowerRoman", start=1)


def configure_chapter_section(section, first_chapter=False):
    """First page of a chapter is centered at the bottom; continuation pages are top-right."""
    _set_section_geometry(section)
    section.different_first_page_header_footer = True
    section.header.is_linked_to_previous = False
    section.footer.is_linked_to_previous = False
    section.first_page_header.is_linked_to_previous = False
    section.first_page_footer.is_linked_to_previous = False
    _clear_header_footer(section.header)
    _clear_header_footer(section.footer)
    _clear_header_footer(section.first_page_header)
    _clear_header_footer(section.first_page_footer)
    add_page_number(section.header.paragraphs[0], WD_ALIGN_PARAGRAPH.RIGHT)
    add_page_number(section.first_page_footer.paragraphs[0], WD_ALIGN_PARAGRAPH.CENTER)
    _set_page_number_type(section, "decimal", start=1 if first_chapter else None)


def configure_document(doc: Document):
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(3)
    section.bottom_margin = Cm(3)
    section.left_margin = Cm(4)
    section.right_margin = Cm(3)
    section.header_distance = Cm(1.25)
    section.footer_distance = Cm(1.25)

    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)
    normal.font.color.rgb = BLACK
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.line_spacing = 1.5
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.first_line_indent = Cm(1.25)

    for name, size, align, before, after in [
        ("Heading 1", 14, WD_ALIGN_PARAGRAPH.CENTER, 18, 4),
        ("Heading 2", 12, WD_ALIGN_PARAGRAPH.LEFT, 8, 4),
        ("Heading 3", 12, WD_ALIGN_PARAGRAPH.LEFT, 8, 4),
    ]:
        style = doc.styles[name]
        style.font.name = "Times New Roman"
        style.font.size = Pt(size)
        style.font.color.rgb = BLACK
        style.font.bold = True
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        style.paragraph_format.alignment = align
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.5
        style.paragraph_format.first_line_indent = Cm(0)

    configure_cover_section(doc.sections[0])


def add_inline(p, text: str):
    # Small Markdown inline subset for readable DOCX output.
    pattern = re.compile(r"(\*\*.*?\*\*|\*.*?\*|`.*?`|\[.*?\]\(.*?\))")
    pos = 0
    for match in pattern.finditer(text):
        if match.start() > pos:
            p.add_run(text[pos:match.start()])
        token = match.group(0)
        if token.startswith("**"):
            run = p.add_run(token[2:-2]); run.bold = True
        elif token.startswith("*"):
            run = p.add_run(token[1:-1]); run.italic = True
        elif token.startswith("`"):
            run = p.add_run(token[1:-1]); run.font.name = "Consolas"; run.font.size = Pt(10)
        else:
            label, url = re.match(r"\[(.*?)\]\((.*?)\)", token).groups()
            run = p.add_run(label); run.underline = True
        pos = match.end()
    if pos < len(text):
        p.add_run(text[pos:])


def add_body(doc, text: str, style: str | None = None, indent: bool = True):
    p = doc.add_paragraph(style=style or "Normal")
    if style is None:
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.line_spacing = 1.5
        p.paragraph_format.first_line_indent = Cm(1.25) if indent else Cm(0)
    add_inline(p, text)
    for r in p.runs:
        r.font.name = "Times New Roman"
        r.font.color.rgb = BLACK
        if r.font.size is None: r.font.size = Pt(12)
    return p


def latex_to_plain(lines: list[str]) -> str:
    """Convert the small LaTeX subset used in the drafts into readable DOCX text."""
    import re
    raw = "\n".join(lines).replace("\\\\", "\n")
    raw = raw.replace(r"\begin{aligned}", "").replace(r"\end{aligned}", "")
    raw = raw.replace(r"\begin{cases}", "").replace(r"\end{cases}", "")
    raw = re.sub(r"\\frac\{([^{}]+)\}\{([^{}]+)\}", r"(\1) / (\2)", raw)
    raw = re.sub(r"\\text\{([^{}]+)\}", r"\1", raw)
    raw = raw.replace(r"\sum", "Σ").replace(r"\pi", "π").replace(r"\Phi", "Φ")
    raw = raw.replace(r"\leq", "≤").replace(r"\geq", "≥").replace(r"\neq", "≠")
    raw = raw.replace(r"\in", "∈").replace(r"\cdot", "×").replace(r"\times", "×")
    raw = raw.replace(r"\quad", "  ").replace(r"\,", " ")
    raw = raw.replace("&", " ").replace("{", "").replace("}", "")
    cleaned = []
    for line in raw.splitlines():
        line = re.sub(r"\s+", " ", line).strip()
        if not line: continue
        if line.startswith("1,") or line.startswith("0,"):
            line = line.replace(", ", ", jika ", 1)
        cleaned.append(line)
    return "\n".join(cleaned)


def add_formula(doc, lines: list[str]):
    plain = latex_to_plain(lines)
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1.0)
    p.paragraph_format.right_indent = Cm(0.5)
    p.paragraph_format.line_spacing = 1.15
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    plain_lines = plain.splitlines()
    for index, line in enumerate(plain_lines):
        run = p.add_run(line + ("\n" if index < len(plain_lines) - 1 else ""))
        run.font.name = "Consolas"
        run.font.size = Pt(10.5)
        run.font.color.rgb = BLACK
    return p


def add_figure(doc: Document, caption: str):
    caption_paragraph = add_body(doc, caption, indent=False)
    caption_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    image_path = FIGURE_IMAGES.get(caption)
    if image_path and image_path.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(str(image_path), width=Cm(13.5))


def parse_table(lines: list[str], doc: Document):
    rows = []
    for line in lines:
        parts = [x.strip() for x in line.strip().strip("|").split("|")]
        if all(re.fullmatch(r":?-{3,}:?", x) for x in parts):
            continue
        rows.append(parts)
    if not rows:
        return
    cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=cols)
    table.style = "Table Grid"
    symbol_col = 1 if rows and len(rows[0]) > 1 and rows[0][1].lower() == 'simbol' else None
    for i, row in enumerate(rows):
        for j in range(cols):
            set_cell_text(table.cell(i, j), row[j] if j < len(row) else "", bold=(i == 0))
            if symbol_col is not None and i > 0 and j == symbol_col:
                table.cell(i, j).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_markdown(doc: Document, path: Path):
    raw = path.read_text(encoding="utf-8").splitlines()
    i = 0
    in_code = False
    code_lines = []
    while i < len(raw):
        line = raw[i].rstrip()
        if line.strip() == r"\[":
            formula_lines = []
            i += 1
            while i < len(raw) and raw[i].strip() != r"\]":
                formula_lines.append(raw[i])
                i += 1
            if i < len(raw): i += 1
            add_formula(doc, formula_lines)
            continue
        if line.startswith("```"): 
            if not in_code:
                in_code = True; code_lines = []
            else:
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Cm(0.5)
                p.paragraph_format.line_spacing = 1
                for n, code in enumerate(code_lines):
                    r = p.add_run(code + ("\n" if n < len(code_lines)-1 else ""))
                    r.font.name = "Consolas"; r.font.size = Pt(9)
                in_code = False
            i += 1; continue
        if in_code:
            code_lines.append(line); i += 1; continue
        if not line.strip() or line.strip() == "---":
            i += 1; continue
        if line.startswith("**Gambar ") and line.endswith("**"):
            caption = line[2:-2].strip()
            add_figure(doc, caption)
            i += 1
            continue
        if line.startswith("|"):
            table_lines = []
            while i < len(raw) and raw[i].strip().startswith("|"):
                table_lines.append(raw[i]); i += 1
            parse_table(table_lines, doc); continue
        if line.startswith("# "):
            title = line[2:].strip()
            if title.startswith("BAB ") or title in {"DAFTAR PUSTAKA", "LAMPIRAN"}:
                section = doc.add_section(WD_SECTION.NEW_PAGE)
                configure_chapter_section(section, first_chapter=(title == "BAB I" or title.startswith("BAB I ")))
            add_body(doc, title, style="Heading 1", indent=False)
        elif line.startswith("## "):
            add_body(doc, line[3:].strip(), style="Heading 2", indent=False)
        elif line.startswith("### "):
            add_body(doc, line[4:].strip(), style="Heading 3", indent=False)
        elif line.startswith("> "):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.6)
            p.paragraph_format.right_indent = Cm(0.3)
            p.paragraph_format.line_spacing = 1.15
            add_inline(p, line[2:].strip())
            for r in p.runs:
                r.italic = True; r.font.name = "Times New Roman"; r.font.size = Pt(11); r.font.color.rgb = BLACK
        elif re.match(r"^\d+\.\s+", line):
            text = re.sub(r"^\d+\.\s+", "", line)
            p = doc.add_paragraph(style="List Number")
            p.paragraph_format.line_spacing = 1.5
            p.paragraph_format.first_line_indent = Cm(0)
            add_inline(p, text)
        elif line.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.line_spacing = 1.5
            p.paragraph_format.first_line_indent = Cm(0)
            add_inline(p, line[2:].strip())
        else:
            add_body(doc, line, indent=True)
        i += 1


def add_title_page(doc: Document):
    for _ in range(5): doc.add_paragraph()
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("DRAFT SKRIPSI"); r.bold = True; r.font.name = "Times New Roman"; r.font.size = Pt(14)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("SISTEM PENDUKUNG KEPUTUSAN PENENTUAN PRIORITAS\nPENGADAAN BARANG DAN JASA PT PRISMA GAPURA\nDENGAN METODE PROMETHEE II BERBASIS WEB")
    r.bold = True; r.font.name = "Times New Roman"; r.font.size = Pt(14)
    for _ in range(5): doc.add_paragraph()
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Disusun sebagai draf kerja penelitian\n\nNama Peneliti\nNIM\n\nProgram Studi\nFakultas\nUniversitas\n2026")
    r.font.name = "Times New Roman"; r.font.size = Pt(12)

    # Front matter mengikuti pola pedoman kampus. Nama dan tanda tangan sengaja
    # dibiarkan sebagai placeholder sampai identitas resmi tersedia.
    front_section = doc.add_section(WD_SECTION.NEW_PAGE)
    configure_front_matter_section(front_section)
    add_body(doc, "LEMBAR PERSETUJUAN", style="Heading 1", indent=False)
    add_body(doc, "Draf skripsi dengan judul:", indent=False)
    p = add_body(doc, "SISTEM PENDUKUNG KEPUTUSAN PENENTUAN PRIORITAS PENGADAAN BARANG DAN JASA PT PRISMA GAPURA DENGAN METODE PROMETHEE II BERBASIS WEB", indent=False)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_body(doc, "diajukan untuk mendapatkan persetujuan pembimbing.", indent=False)
    for _ in range(5): doc.add_paragraph()
    t = doc.add_table(rows=4, cols=2); t.style = "Table Grid"
    for i, (a, b) in enumerate([("Nama Peneliti", "................................................"), ("NIM", "................................................"), ("Pembimbing", "................................................"), ("Tanggal", "................................................")]):
        set_cell_text(t.cell(i,0), a); set_cell_text(t.cell(i,1), b)

    doc.add_page_break()
    add_body(doc, "LEMBAR PENGESAHAN", style="Heading 1", indent=False)
    add_body(doc, "Draf ini telah diperiksa dan dapat digunakan sebagai bahan review penelitian.", indent=False)
    for _ in range(5): doc.add_paragraph()
    t = doc.add_table(rows=3, cols=3); t.style = "Table Grid"
    for j, h in enumerate(["Jabatan", "Nama", "Tanda Tangan"]): set_cell_text(t.cell(0,j), h, bold=True)
    for i, row in enumerate([("Pembimbing", "................................", "................................"), ("Ketua Program Studi", "................................", "................................")], start=1):
        for j, value in enumerate(row): set_cell_text(t.cell(i,j), value)

    doc.add_page_break()
    add_body(doc, "PERNYATAAN KEASLIAN", style="Heading 1", indent=False)
    add_body(doc, "Saya yang bertanda tangan di bawah ini menyatakan bahwa draf skripsi ini disusun oleh peneliti sendiri. Setiap teori, data, dan hasil penelitian yang berasal dari sumber lain akan dicantumkan sesuai ketentuan sitasi yang berlaku. Data pada draf ini masih berstatus simulasi dan akan diperbarui sesuai izin penggunaan data penelitian.", indent=True)
    for _ in range(6): doc.add_paragraph()
    add_body(doc, "Nama Peneliti: ................................................", indent=False)
    add_body(doc, "NIM: ................................................", indent=False)
    add_body(doc, "Tanda tangan: ................................................", indent=False)

    doc.add_page_break()
    add_body(doc, "ABSTRAK", style="Heading 1", indent=False)
    add_body(doc, "Penelitian ini bertujuan membangun Sistem Pendukung Keputusan untuk menentukan prioritas pengadaan barang dan jasa PT Prisma Gapura menggunakan metode PROMETHEE II berbasis web. Alternatif yang digunakan terdiri atas barang dan jasa proteksi APAR. Penilaian dilakukan menggunakan lima kriteria, yaitu C1 Urgensi Kebutuhan, C2 Frekuensi Transaksi, C3 Total Biaya Pengadaan, C4 Waktu Pengadaan, dan C5 Tingkat Kekritisan. C1, C2, dan C5 merupakan kriteria benefit, sedangkan C3 dan C4 merupakan kriteria cost.", indent=True)
    add_body(doc, "Fungsi preferensi yang digunakan pada seluruh kriteria adalah Usual atau Type 1. Sistem melakukan perbandingan berpasangan, menghitung indeks preferensi, leaving flow, entering flow, dan net flow untuk menghasilkan ranking. Data yang digunakan pada draf ini merupakan data simulasi berbasis pola operasional dan belum menjadi data penelitian final. Hasil implementasi awal menunjukkan bahwa sistem dapat mengelola data, bobot, perhitungan, dan hasil ranking melalui akses Admin.", indent=True)
    add_body(doc, "Kata kunci: Sistem Pendukung Keputusan, PROMETHEE II, prioritas pengadaan, barang dan jasa, APAR.", indent=False)

    doc.add_page_break()
    add_body(doc, "ABSTRACT", style="Heading 1", indent=False)
    add_body(doc, "This study aims to develop a web-based Decision Support System for prioritizing goods and services procurement at PT Prisma Gapura using the PROMETHEE II method. The alternatives consist of fire protection goods and services. The evaluation uses five criteria: C1 Requirement Urgency, C2 Transaction Frequency, C3 Total Procurement Cost, C4 Procurement Lead Time, and C5 Criticality Level. C1, C2, and C5 are benefit criteria, while C3 and C4 are cost criteria.", indent=True)
    add_body(doc, "The Usual or Type 1 preference function is applied to all criteria. The system performs pairwise comparisons, calculates preference indices, leaving flow, entering flow, and net flow to produce a ranking. The dataset in this draft is a simulated dataset based on operational patterns and is not the final research dataset. The initial implementation supports data management, weight adjustment, calculation, and ranking through an Admin account.", indent=True)
    add_body(doc, "Keywords: Decision Support System, PROMETHEE II, procurement priority, goods and services, fire extinguisher.", indent=False)

    doc.add_page_break()
    add_body(doc, "KATA PENGANTAR", style="Heading 1", indent=False)
    add_body(doc, "Puji dan syukur penulis panjatkan kepada Tuhan Yang Maha Esa karena draf penelitian ini dapat disusun sebagai bahan pengembangan skripsi. Draf ini membahas perancangan Sistem Pendukung Keputusan penentuan prioritas pengadaan barang dan jasa PT Prisma Gapura menggunakan metode PROMETHEE II berbasis web.", indent=True)
    add_body(doc, "Penyusunan draf ini masih memerlukan validasi data, arahan pembimbing, dan penyesuaian terhadap pedoman akademik yang berlaku. Penulis menyadari bahwa isi draf ini masih memiliki kekurangan, terutama karena data yang digunakan masih berupa data simulasi. Oleh sebab itu, kritik dan saran diperlukan untuk penyempurnaan penelitian.", indent=True)
    add_body(doc, "Penulis mengucapkan terima kasih kepada pihak-pihak yang memberikan bimbingan, masukan, dan dukungan selama proses penyusunan penelitian ini. Semoga penelitian ini dapat memberikan manfaat bagi pengembangan sistem pendukung keputusan dan pengelolaan prioritas pengadaan.", indent=True)
    for _ in range(4): doc.add_paragraph()
    add_body(doc, "Penulis", indent=False)

    doc.add_page_break()
    add_body(doc, "DAFTAR ISI", style="Heading 1", indent=False)
    toc_items = [
        ("ABSTRAK", 0), ("ABSTRACT", 0), ("KATA PENGANTAR", 0), ("DAFTAR ISI", 0),
        ("DAFTAR GAMBAR", 0), ("DAFTAR TABEL", 0), ("DAFTAR LAMPIRAN", 0),
        ("BAB I PENDAHULUAN", 0),
        ("1.1 Latar Belakang", 1), ("1.2 Identifikasi Masalah", 1), ("1.3 Rumusan Masalah", 1),
        ("1.4 Batasan Penelitian", 1), ("1.5 Tujuan Penelitian", 1), ("1.6 Manfaat Penelitian", 1),
        ("1.6.1 Manfaat Teoretis", 2), ("1.6.2 Manfaat Praktis", 2), ("1.7 Metodologi Penelitian", 1),
        ("1.7.1 Metode Pengumpulan Data", 2), ("1.7.2 Metode Pengembangan Sistem", 2), ("1.7.3 Metode Pengujian", 2), ("1.8 Sistematika Penulisan", 1),
        ("BAB II LANDASAN TEORI", 0), ("2.1 Penelitian Terdahulu", 1), ("2.2 Tinjauan Pustaka", 1),
        ("2.2.1 Sistem Informasi", 2), ("2.2.2 Sistem Pendukung Keputusan", 2), ("2.2.3 Pengadaan Barang dan Jasa", 2),
        ("2.2.4 Prioritas Penanganan Pengadaan", 2), ("2.2.5 Website", 2), ("2.2.6 Metode PROMETHEE II", 2), ("2.2.7 Fungsi Preferensi Usual", 2), ("2.2.8 Benefit dan Cost", 2), ("2.2.9 Bobot Kriteria", 2), ("2.2.10 Indeks Preferensi", 2), ("2.2.11 Leaving Flow, Entering Flow, dan Net Flow", 2),
        ("2.3 Perangkat Lunak Pendukung", 1), ("2.3.1 Python", 2), ("2.3.2 MySQL", 2), ("2.3.3 HTML5", 2),
        ("2.3.4 CSS3", 2), ("2.3.5 JavaScript", 2), ("2.3.6 SQLite", 2), ("2.3.7 Python Standard Library", 2),
        ("2.3.8 MySQL Connector/Python", 2), ("2.3.9 Visual Studio Code", 2), ("2.3.10 XAMPP", 2), ("2.3.11 draw.io", 2),
        ("2.4 Unified Modeling Language (UML)", 1), ("2.4.1 Use Case Diagram", 2), ("2.4.2 Activity Diagram", 2),
        ("2.4.3 Sequence Diagram", 2), ("2.4.4 Class Diagram", 2), ("2.5 Perancangan Basis Data", 1),
        ("2.5.1 Basis Data", 2), ("2.5.2 Entity Relationship Diagram (ERD)", 2), ("2.5.3 Transformasi ERD ke LRS", 2),
        ("2.5.4 Logical Record Structure (LRS)", 2), ("2.6 Pengujian Sistem", 1), ("2.6.1 Pengujian Blackbox", 2),
        ("2.6.2 Verifikasi Numerik", 2), ("2.7 Kerangka Pemikiran", 1),
        ("BAB III ANALISA DAN PERANCANGAN", 0), ("3.1 Analisa Sistem", 1), ("3.1.1 Analisa Sistem Berjalan", 2),
        ("3.1.2 Analisa Sistem Usulan", 2), ("3.1.2.1 Batasan Sistem Usulan", 3), ("3.1.3 Kebutuhan Fungsional", 2), ("3.1.4 Kebutuhan Nonfungsional", 2), ("3.1.5 Data Penelitian", 2), ("3.2 Penerapan Metode PROMETHEE II", 1),
        ("3.2.1 Urutan Tahapan Perhitungan", 2), ("3.2.2 Penentuan Alternatif", 2), ("3.2.3 Penentuan Kriteria dan Arah Preferensi", 2), ("3.2.4 Penentuan dan Normalisasi Bobot", 2), ("3.2.5 Pembentukan Matriks Evaluasi", 2), ("3.2.6 Perbandingan Berpasangan dan Fungsi Preferensi Usual", 2), ("3.2.7 Contoh Perhitungan Nilai Preferensi", 2), ("3.2.8 Perhitungan Indeks Preferensi Multikriteria", 2), ("3.2.9 Perhitungan Leaving Flow dan Entering Flow", 2), ("3.2.10 Penyusunan Ranking dan Interpretasi Output", 2), ("3.2.11 Alur Algoritmik Sistem", 2), ("3.2.12 Pemeriksaan Konsistensi Perhitungan", 2), ("3.3 Perancangan Basis Data", 1), ("3.3.1 ERD", 2),
        ("3.3.2 Transformasi ERD ke LRS", 2), ("3.3.3 LRS", 2), ("3.3.4 Spesifikasi Tabel Users dan Criteria", 2), ("3.3.5 Spesifikasi Tabel Alternatives dan Transactions", 2), ("3.3.6 Spesifikasi Tabel Settings dan Contoh Record", 2), ("3.3.7 Hubungan Data dan Proses CRUD", 2),
        ("3.4 Perancangan UML", 1), ("3.4.1 Use Case Diagram", 2), ("3.4.2 Activity Diagram", 2),
        ("3.4.3 Sequence Diagram", 2), ("3.4.4 Class Diagram", 2), ("3.4.5 Validasi Rancangan UML terhadap Fitur", 2), ("3.5 Perancangan Antarmuka", 1),
        ("3.5.1 Rancangan Landing Page Profile", 2), ("3.5.2 Rancangan Login Admin", 2), ("3.5.3 Rancangan Dashboard Admin", 2),
        ("3.5.4 Rancangan Data Transaksi", 2), ("3.5.5 Rancangan Kriteria dan Bobot", 2), ("3.5.6 Rancangan Kelola Data", 2),
        ("3.5.7 Rancangan Perhitungan PROMETHEE II", 2), ("3.5.8 Rancangan Hasil Ranking", 2), ("3.5.9 Rancangan Panduan", 2), ("3.5.10 Prinsip Keterhubungan Rancangan dan Implementasi", 2),
        ("BAB IV IMPLEMENTASI DAN PENGUJIAN", 0), ("4.1 Spesifikasi", 1), ("4.1.1 Perangkat Lunak", 2),
        ("4.1.2 Perangkat Keras", 2), ("4.2 Implementasi Program", 1), ("4.2.1 Landing Page Profile", 2), ("4.2.2 Login Admin", 2),
        ("4.2.3 Dashboard Admin", 2), ("4.2.4 Data Transaksi", 2), ("4.2.5 Kelola Data", 2),
        ("4.2.6 Kriteria dan Bobot", 2), ("4.2.7 Perhitungan PROMETHEE II", 2), ("4.2.8 Hasil Ranking", 2),
        ("4.3 Pengujian Sistem", 1), ("4.3.1 Pengujian Blackbox", 2), ("4.3.2 Verifikasi Numerik", 2),
        ("4.4 Analisis Hasil Ranking", 1), ("4.5 Pembahasan Hasil Pengujian", 1), ("BAB V PENUTUP", 0), ("5.1 Kesimpulan", 1), ("5.2 Saran", 1),
        ("DAFTAR PUSTAKA", 0), ("LAMPIRAN", 0),
    ]
    for item, level in toc_items:
        p = add_body(doc, item, indent=False)
        p.paragraph_format.left_indent = Cm(0.4 * level)
    add_body(doc, "Catatan: daftar isi otomatis perlu diperbarui melalui fitur References → Table of Contents setelah struktur dan nomor halaman final.", indent=False)

    doc.add_page_break()
    add_body(doc, "DAFTAR GAMBAR", style="Heading 1", indent=False)
    for item in [
        "Gambar 3.1 Arsitektur Sistem SPK", "Gambar 3.2 Entity Relationship Diagram (ERD)",
        "Gambar 3.3 Logical Record Structure (LRS)", "Gambar 3.4 Use Case Diagram",
        "Gambar 3.5 Activity Diagram Perhitungan PROMETHEE II", "Gambar 3.6 Sequence Diagram Proses Perhitungan",
        "Gambar 3.7 Class Diagram", "Gambar 3.8 Rancangan Landing Page Profile", "Gambar 3.9 Rancangan Login Admin",
        "Gambar 3.10 Rancangan Dashboard Admin", "Gambar 3.11 Rancangan Data Transaksi", "Gambar 3.12 Rancangan Kriteria dan Bobot",
        "Gambar 3.13 Rancangan Kelola Data", "Gambar 3.14 Rancangan Perhitungan PROMETHEE II", "Gambar 3.15 Rancangan Hasil Ranking",
        "Gambar 3.16 Rancangan Panduan", "Gambar 4.1 Implementasi Landing Page Profile", "Gambar 4.2 Implementasi Halaman Login Admin",
        "Gambar 4.3 Implementasi Dashboard Admin", "Gambar 4.4 Implementasi Halaman Data Transaksi", "Gambar 4.5 Implementasi Halaman Kelola Data",
        "Gambar 4.6 Implementasi Halaman Kriteria dan Bobot", "Gambar 4.7 Implementasi Halaman Perhitungan PROMETHEE II", "Gambar 4.8 Implementasi Halaman Hasil Ranking",
    ]:
        add_body(doc, item, indent=False)

    doc.add_page_break()
    add_body(doc, "DAFTAR TABEL", style="Heading 1", indent=False)
    for item in [
        "Tabel 2.1 Ringkasan Penelitian Terdahulu", "Tabel 2.2 Bobot Kriteria Penelitian",
        "Tabel 2.3 Simbol-Simbol Use Case Diagram", "Tabel 2.4 Simbol-Simbol Activity Diagram",
        "Tabel 2.5 Simbol-Simbol Sequence Diagram", "Tabel 2.6 Simbol-Simbol Class Diagram",
        "Tabel 2.7 Simbol-Simbol Entity Relationship Diagram (ERD)",
        "Tabel 3.1 Tahapan Sistem Berjalan", "Tabel 3.2 Alur Sistem Usulan", "Tabel 3.3 Kebutuhan Fungsional",
        "Tabel 3.4 Kebutuhan Nonfungsional", "Tabel 3.5 Ringkasan Data Penelitian", "Tabel 3.6 Matriks Evaluasi Alternatif Penelitian",
        "Tabel 3.7 Contoh Data Transaksi", "Tabel 3.8 Aturan Pembentukan Nilai Evaluasi", "Tabel 3.9 Tahapan Perhitungan PROMETHEE II pada Sistem",
        "Tabel 3.10 Kriteria dan Arah Preferensi", "Tabel 3.11 Bobot Awal dan Bobot Ternormalisasi", "Tabel 3.12 Perbandingan A01 terhadap A02 dengan Fungsi Usual",
        "Tabel 3.13 Kontribusi Bobot pada Indeks Preferensi A01 dan A02", "Tabel 3.14 Contoh Hasil Flow Alternatif Teratas", "Tabel 3.15 Hasil Ranking Dataset Simulasi",
        "Tabel 3.16 Pemeriksaan Konsistensi Metode", "Tabel 3.17 Spesifikasi Tabel users", "Tabel 3.18 Spesifikasi Tabel criteria",
        "Tabel 3.19 Spesifikasi Tabel alternatives", "Tabel 3.20 Spesifikasi Tabel transactions", "Tabel 3.21 Spesifikasi Tabel settings",
        "Tabel 3.22 Contoh Record Database", "Tabel 3.23 Hubungan Antarentitas", "Tabel 3.24 Pemetaan CRUD terhadap Data",
        "Tabel 3.25 Deskripsi Use Case Admin", "Tabel 3.26 Rincian Aktivitas Utama", "Tabel 3.27 Urutan Pesan pada Sequence Diagram",
        "Tabel 3.28 Deskripsi Kelas Sistem", "Tabel 3.29 Pemetaan Kebutuhan ke Rancangan", "Tabel 3.30 Pemetaan Mockup terhadap Kebutuhan Sistem",
        "Tabel 4.1 Spesifikasi Perangkat Lunak", "Tabel 4.2 Spesifikasi Perangkat Keras", "Tabel 4.3 Contoh Pengujian",
        "Tabel 4.4 Rekapitulasi Pengujian Black Box",
    ]:
        add_body(doc, item, indent=False)

    doc.add_page_break()
    add_body(doc, "DAFTAR LAMPIRAN", style="Heading 1", indent=False)
    for item in [
        "Lampiran 1 Matriks Evaluasi", "Lampiran 2 Matriks Preferensi Usual",
        "Lampiran 3 Hasil Flow dan Ranking", "Lampiran 4 Perhitungan PROMETHEE II",
        "Lampiran 5 Hasil Pengujian Sistem", "Lampiran 6 LRS dan Spesifikasi Basis Data",
        "Lampiran 7 Diagram Perancangan", "Lampiran 8 Screenshot Sistem",
        "Lampiran 9 Catatan Data Simulasi Jasa", "Lampiran 10 Referensi PDF", "Lampiran 11 Validasi Data dan Bobot",
        "Lampiran 12 Verifikasi Numerik", "Lampiran 13 Hasil Pengujian Lengkap", "Lampiran 14 Validasi UML dan Database",
        "Lampiran 15 Hasil Uji MySQL", "Lampiran 16 Audit Sitasi dan Download PDF",
    ]:
        add_body(doc, item, indent=False)


def main():
    doc = Document()
    configure_document(doc)
    add_title_page(doc)
    for chapter in CHAPTERS:
        add_markdown(doc, chapter)
    doc.save(OUT)
    print(f"Wrote {OUT}")


if __name__ == "__main__":
    main()
