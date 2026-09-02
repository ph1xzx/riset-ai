import sys
from docx import Document

def first_long_body(d):
    for p in d.paragraphs:
        st = p.style.name if p.style else ""
        if st and st.startswith("Heading"): continue
        if len(p.text) > 80: return p
    return None

def facts(path):
    d = Document(path)
    sec = d.sections[0]
    margins = [round(x/360000, 2) if x else None for x in
               (sec.top_margin, sec.right_margin, sec.bottom_margin, sec.left_margin)]
    bp = first_long_body(d)
    font = size = spacing = None
    if bp is not None:
        spacing = bp.paragraph_format.line_spacing
        r = bp.runs[0] if bp.runs else None
        if r is not None:
            font = r.font.name; size = r.font.size.pt if r.font.size else None
    heads = [( (p.style.name if p.style else "?"), p.text.strip()[:60]) for p in d.paragraphs
             if p.style and p.style.name and p.style.name.startswith("Heading") and p.text.strip()]
    imgs = len(d.inline_shapes)
    tables = len(d.tables)
    words = sum(len(p.text.split()) for p in d.paragraphs if p.text)
    return dict(margins=margins, font=font, size=size, spacing=spacing,
                heads=heads, imgs=imgs, tables=tables, words=words)

a, b = facts(sys.argv[1]), facts(sys.argv[2])
def show(tag, f):
    print(f"== {tag} == margins cm {f['margins']} | body font {f['font']} {f['size']}pt | spasi {f['spacing']}")
    print(f"   headings {len(f['heads'])} | images {f['imgs']} | tables {f['tables']} | words {f['words']}")
show("ORIGINAL", a); show("EXPORT  ", b)
at = [t for _, t in a["heads"]]; bt = [t for _, t in b["heads"]]
missing = [t for t in at if t not in bt]
extra = [t for t in bt if t not in at]
print("heading hilang di export:", len(missing), "| heading tambahan:", len(extra))
for t in missing[:8]: print("  hilang:", t)
for t in extra[:8]: print("  tambah:", t)
# cek kebocoran tag / kata dempet
d = Document(sys.argv[2])
leaks = []
for p in d.paragraphs:
    if "<" in p.text and ("strong" in p.text or "<p" in p.text or "<h" in p.text): leaks.append(p.text[:60])
bad = [w for w in ("PRIORITASPENGADAAN", "GAPURADENGAN", "penelitianNama") if w in "".join(p.text for p in d.paragraphs)]
print("literal-tag leaks:", len(leaks), leaks[:3])
print("kata dempet ditemukan:", bad if bad else "tidak ada")
